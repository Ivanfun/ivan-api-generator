import pandas as pd
from docx import Document
from docx.shared import Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

def load_sql_properties(path):
    sql_dict = {}
    with open(path, 'r', encoding='utf-8') as f:
        for line in f:
            if '=' in line and not line.startswith('#'):
                key, val = line.split('=', 1)
                sql_dict[key.strip()] = val.strip()
    return sql_dict

def generate_api_doc(excel_path: Path, word_template_path: Path, output_path: Path, sql_properties_path: Path):
    sheets = pd.read_excel(excel_path, sheet_name=None)
    api_hierarchy_df = sheets.get('API階層表')
    if api_hierarchy_df is None:
        raise ValueError("Excel 中找不到 'API階層表' 工作表。")

    api_data = {}
    for name, df in sheets.items():
        if 'API代碼' not in df.columns:
            continue
        for api in df['API代碼'].dropna().unique():
            if api not in api_data:
                api_data[api] = {}
            api_data[api][name] = df[df['API代碼'] == api]

    sql_map = load_sql_properties(sql_properties_path)
    doc = Document(word_template_path)

    batch_list = api_hierarchy_df[['批次代碼', '批次說明']].drop_duplicates()
    for row in batch_list.itertuples(index=False):
        batch_code, batch_desc = row
        group_df = api_hierarchy_df[
            (api_hierarchy_df['批次代碼'] == batch_code) &
            (api_hierarchy_df['批次說明'] == batch_desc)
        ]

        doc.add_paragraph(f'{batch_code} ({batch_desc})', style='Heading 2')
        api_list_table = doc.add_table(rows=1, cols=3)
        api_list_table.style = 'Table Grid'
        api_list_table.autofit = False
        headers = ['順序', 'API代碼', 'API說明']
        widths_cm = [1.24, 7, 10.79]

        for i in range(3):
            cell = api_list_table.cell(0, i)
            cell.text = headers[i]
            cell.width = Cm(widths_cm[i])
            tcPr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), "D9D9D9")
            tcPr.append(shd)

        for _, api_row in group_df.iterrows():
            cells = api_list_table.add_row().cells
            cells[0].text = str(int(api_row['API順序']))
            cells[1].text = str(api_row['API代碼'])
            cells[2].text = str(api_row['API說明'])
            for i in range(3):
                cells[i].width = Cm(widths_cm[i])

        doc.add_paragraph('')

        for _, api_row in group_df.iterrows():
            api_code = api_row['API代碼']
            doc.add_paragraph(api_code, style='Heading 4')
            if api_code not in api_data:
                continue

            if 'API清單' in api_data[api_code]:
                api_df = api_data[api_code]['API清單']
                if not api_df.empty:
                    api_detail_table = doc.add_table(rows=1, cols=3)
                    api_detail_table.style = 'Table Grid'
                    api_detail_table.autofit = False
                    headers = ['序', '參數', '設定值']
                    widths_cm = [1.24, 3.5, 14.29]

                    for i in range(3):
                        cell = api_detail_table.cell(0, i)
                        cell.text = headers[i]
                        cell.width = Cm(widths_cm[i])
                        tcPr = cell._tc.get_or_add_tcPr()
                        shd = OxmlElement('w:shd')
                        shd.set(qn('w:fill'), "D9D9D9")
                        tcPr.append(shd)

                    param_names = [
                        'API代碼', 'API簡述', 'API說明', 'API行為類型', '資料庫連線名稱',
                        '執行類型', '語法設定鍵值', '驗證金鑰', '是否編碼', '語法'
                    ]
                    for i, param in enumerate(param_names, start=1):
                        row_cells = api_detail_table.add_row().cells
                        row_cells[0].text = str(i)
                        row_cells[1].text = param
                        if param == '語法':
                            config_key = api_df['語法設定鍵值'].values[0]
                            raw_value = sql_map.get(config_key, '🔍 查無對應 SQL')
                            value = raw_value.replace('\\n', '\n').replace('\\t', '\t').replace('\\=', '=')
                        else:
                            value = api_df[param].values[0] if param in api_df.columns else ''
                        row_cells[2].text = str(value)
                        for j in range(3):
                            row_cells[j].width = Cm(widths_cm[j])
                        for j in [0, 1]:
                            tcPr = row_cells[j]._tc.get_or_add_tcPr()
                            shd = OxmlElement('w:shd')
                            shd.set(qn('w:fill'), "F2F2F2")
                            tcPr.append(shd)

                    doc.add_paragraph('')

            section_order = [
                ('參數驗證', ['序', '屬性名', '預設值', '說明'], [1.24, 3.5, 8.89, 5.4]),
                ('WebService', ['序', '主機代碼', '主機名稱', '主機IP', '啟用'], [1.24, 3.5, 4.57, 4.32, 5.4]),
                ('IP權限設定', ['IP', '說明'], [8.74, 10.25]),
                ('輸出設定', ['節點階層', '父階層關聯鍵值', '子階層關聯鍵值', '輸出參數'], [3.24, 4.5, 5.89, 5.4])
            ]
            for sheet_name, headers, widths_cm in section_order:
                if sheet_name in api_data[api_code]:
                    df = api_data[api_code][sheet_name]
                    if not df.empty:
                        doc.add_paragraph(sheet_name, style='Heading 5')
                        table = doc.add_table(rows=1, cols=len(headers))
                        table.style = 'Table Grid'
                        table.autofit = False
                        for i, header in enumerate(headers):
                            cell = table.cell(0, i)
                            cell.text = header
                            cell.width = Cm(widths_cm[i])
                            tcPr = cell._tc.get_or_add_tcPr()
                            shd = OxmlElement('w:shd')
                            shd.set(qn('w:fill'), "D9D9D9")
                            tcPr.append(shd)
                        for _, row_data in df.iterrows():
                            row_cells = table.add_row().cells
                            for j, header in enumerate(headers):
                                value = row_data.get(header, '')
                                if header in ['序', '節點階層'] and pd.notnull(value):
                                    try:
                                        value = int(value) if float(value).is_integer() else value
                                    except ValueError:
                                        pass
                                row_cells[j].text = str(value)
                                row_cells[j].width = Cm(widths_cm[j])
                        doc.add_paragraph('')

    doc.save(output_path)

